"""
Unit tests for Document Extraction Pipeline [CARD-145].
[REQ-DOC-001] [REQ-DOC-002] [REQ-DOC-003] [REQ-DOC-004]
"""

from pathlib import Path
import pytest
from pypdf import PdfWriter
import openpyxl
from docx import Document

from src.application.skills.document_extractors import (
    extract_csv_table,
    extract_excel_tables,
    extract_docx_text,
    extract_pdf_text,
    extract_document,
)
from src.application.skills.document_tools import read_document_file


def test_extract_csv_table(tmp_path):
    csv_file = tmp_path / "data.csv"
    csv_file.write_text("Name,Role,Score\nAlice,Admin,95\nBob,Engineer,88\n", encoding="utf-8")

    result = extract_csv_table(str(csv_file))
    assert "| Name | Role | Score |" in result
    assert "| Alice | Admin | 95 |" in result
    assert "| Bob | Engineer | 88 |" in result


def test_extract_excel_tables(tmp_path):
    xlsx_file = tmp_path / "report.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Financials"
    ws.append(["Quarter", "Revenue", "Profit"])
    ws.append(["Q1", "$100k", "$25k"])
    ws.append(["Q2", "$120k", "$30k"])
    wb.save(xlsx_file)

    result = extract_excel_tables(str(xlsx_file))
    assert "### Sheet: Financials" in result
    assert "| Quarter | Revenue | Profit |" in result
    assert "| Q1 | $100k | $25k |" in result
    assert "| Q2 | $120k | $30k |" in result


def test_extract_docx_text(tmp_path):
    docx_file = tmp_path / "memo.docx"
    doc = Document()
    doc.add_heading("Q3 Strategic Goals", level=1)
    doc.add_paragraph("First milestone is deploying AutoReiv 2.0 to production.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Target"
    table.cell(1, 0).text = "Uptime"
    table.cell(1, 1).text = "99.9%"
    doc.save(docx_file)

    result = extract_docx_text(str(docx_file))
    assert "# Q3 Strategic Goals" in result
    assert "First milestone is deploying AutoReiv 2.0 to production." in result
    assert "| Metric | Target |" in result
    assert "| Uptime | 99.9% |" in result


def test_extract_pdf_text(tmp_path):
    pdf_file = tmp_path / "doc.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    writer.write(pdf_file)

    result = extract_pdf_text(str(pdf_file))
    assert "--- [Page 1" in result or "Total Pages: 1" in result


def test_extract_document_dispatcher(tmp_path):
    csv_file = tmp_path / "notes.txt"
    csv_file.write_text("Hello world text file.", encoding="utf-8")

    res = extract_document(str(csv_file))
    assert res["success"] is True
    assert "Hello world text file." in res["content"]


def test_read_document_file_tool(tmp_path):
    csv_file = tmp_path / "stats.csv"
    csv_file.write_text("Year,Users\n2025,500\n2026,2000\n", encoding="utf-8")

    out = read_document_file(str(csv_file))
    assert "2025" in out
    assert "2000" in out

    missing_out = read_document_file(str(tmp_path / "does_not_exist.pdf"))
    assert "Error:" in missing_out
