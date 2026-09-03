"""
Document extractors for PDF, Excel, Word, CSV, and text documents [CARD-145].
[REQ-DOC-001] [REQ-DOC-002] [REQ-DOC-003]
"""

import csv
import logging
from pathlib import Path
from typing import Any, Dict, List

logger = logging.getLogger(__name__)


def _format_markdown_table(headers: List[str], rows: List[List[str]]) -> str:
    """Format a list of header strings and row lists into a clean markdown table."""
    if not headers and not rows:
        return ""
    if not headers and rows:
        headers = [f"Col {i + 1}" for i in range(len(rows[0]))]

    clean_headers = [str(h or "").replace("\n", " ").replace("|", "\\|").strip() for h in headers]
    lines = [
        "| " + " | ".join(clean_headers) + " |",
        "| " + " | ".join(["---"] * len(clean_headers)) + " |",
    ]
    for r in rows:
        padded_row = list(r) + [""] * max(0, len(clean_headers) - len(r))
        clean_row = [str(c or "").replace("\n", " ").replace("|", "\\|").strip() for c in padded_row[:len(clean_headers)]]
        lines.append("| " + " | ".join(clean_row) + " |")
    return "\n".join(lines)


def extract_csv_table(path: str, max_rows: int = 100) -> str:
    """[REQ-DOC-002] Parse CSV file into a formatted markdown table."""
    p = Path(path)
    if not p.exists() or not p.is_file():
        raise FileNotFoundError(f"File not found: {path}")

    with p.open("r", encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f)
        all_rows = list(reader)

    if not all_rows:
        return "(Empty CSV file)"

    headers = all_rows[0]
    data_rows = all_rows[1:max_rows + 1]
    truncated_msg = f"\n*(Showing first {len(data_rows)} of {len(all_rows) - 1} data rows)*" if len(all_rows) - 1 > max_rows else ""
    return _format_markdown_table(headers, data_rows) + truncated_msg


def extract_excel_tables(path: str, max_rows_per_sheet: int = 50) -> str:
    """[REQ-DOC-002] Parse Excel (.xlsx, .xlsm) into formatted markdown tables per worksheet."""
    p = Path(path)
    if not p.exists() or not p.is_file():
        raise FileNotFoundError(f"File not found: {path}")

    try:
        import openpyxl
    except ImportError:
        return "(openpyxl is not installed, cannot extract Excel tables)"

    wb = openpyxl.load_workbook(p, data_only=True, read_only=True)
    sheet_blocks = []

    for sheetname in wb.sheetnames:
        ws = wb[sheetname]
        raw_rows = []
        for row in ws.iter_rows(values_only=True):
            if any(cell is not None for cell in row):
                raw_rows.append([str(c) if c is not None else "" for c in row])
            if len(raw_rows) > max_rows_per_sheet:
                break

        if not raw_rows:
            sheet_blocks.append(f"### Sheet: {sheetname}\n*(Empty worksheet)*")
            continue

        headers = raw_rows[0]
        data_rows = raw_rows[1:max_rows_per_sheet + 1]
        table_md = _format_markdown_table(headers, data_rows)
        sheet_blocks.append(f"### Sheet: {sheetname}\n{table_md}")

    wb.close()
    return "\n\n".join(sheet_blocks) if sheet_blocks else "(Workbook contains no sheets)"


def extract_docx_text(path: str) -> str:
    """[REQ-DOC-003] Parse Word document (.docx) into structured markdown text and tables."""
    p = Path(path)
    if not p.exists() or not p.is_file():
        raise FileNotFoundError(f"File not found: {path}")

    try:
        import docx
    except ImportError:
        return "(python-docx is not installed, cannot extract Word document)"

    doc = docx.Document(p)
    blocks = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name.lower() if para.style else ""
        if "heading 1" in style_name:
            blocks.append(f"# {text}")
        elif "heading 2" in style_name:
            blocks.append(f"## {text}")
        elif "heading 3" in style_name:
            blocks.append(f"### {text}")
        else:
            blocks.append(text)

    for table in doc.tables:
        t_rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                t_rows.append(cells)
        if t_rows:
            headers = t_rows[0]
            data = t_rows[1:]
            blocks.append(_format_markdown_table(headers, data))

    return "\n\n".join(blocks) if blocks else "(Empty Word document)"


def extract_pdf_text(path: str, max_pages: int = 20) -> str:
    """[REQ-DOC-001] Parse PDF into clean text with page delimiters."""
    p = Path(path)
    if not p.exists() or not p.is_file():
        raise FileNotFoundError(f"File not found: {path}")

    try:
        import pypdf
    except ImportError:
        return "(pypdf is not installed, cannot extract PDF)"

    reader = pypdf.PdfReader(p)
    total_pages = len(reader.pages)
    if total_pages == 0:
        return "(Empty PDF document)"

    pages_to_read = min(total_pages, max_pages)
    page_blocks = []

    for i in range(pages_to_read):
        page = reader.pages[i]
        text = page.extract_text() or ""
        trimmed = text.strip()
        header = f"--- [Page {i + 1} of {total_pages}] ---"
        if trimmed:
            page_blocks.append(f"{header}\n{trimmed}")
        else:
            page_blocks.append(f"{header}\n*(No text on this page or scanned image)*")

    if total_pages > max_pages:
        page_blocks.append(f"\n*(Note: Showing first {max_pages} of {total_pages} pages)*")

    return f"**Total Pages: {total_pages}**\n\n" + "\n\n".join(page_blocks)


def extract_document(path: str, max_pages: int = 20, max_rows: int = 50) -> Dict[str, Any]:
    """Universal dispatcher to extract text and tables from any supported document format."""
    p = Path(path)
    if not p.exists() or not p.is_file():
        return {
            "success": False,
            "error": f"File not found: {path}",
            "content": "",
            "format": "unknown",
        }

    ext = p.suffix.lower()
    try:
        if ext == ".pdf":
            content = extract_pdf_text(str(p), max_pages=max_pages)
            fmt = "pdf"
        elif ext in (".xlsx", ".xlsm", ".xls"):
            content = extract_excel_tables(str(p), max_rows_per_sheet=max_rows)
            fmt = "excel"
        elif ext == ".csv":
            content = extract_csv_table(str(p), max_rows=max_rows)
            fmt = "csv"
        elif ext in (".docx", ".doc"):
            content = extract_docx_text(str(p))
            fmt = "docx"
        elif ext in (".txt", ".md", ".json", ".yaml", ".yml", ".py", ".sql", ".sh", ".log", ".html", ".xml"):
            content = p.read_text(encoding="utf-8", errors="replace")
            fmt = "text"
        else:
            return {
                "success": False,
                "error": f"Unsupported document format: {ext}",
                "content": "",
                "format": ext,
            }

        return {
            "success": True,
            "format": fmt,
            "content": content,
            "file_size": p.stat().st_size,
            "filename": p.name,
        }
    except Exception as e:
        logger.exception("Failed to extract document: %s", path)
        return {
            "success": False,
            "error": str(e),
            "content": "",
            "format": ext,
        }
